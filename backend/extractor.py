"""
Compliance Pilot — Question extractor.
Uses LLM (Claude Code) as primary extractor, falls back to Python heuristics.
"""

import re
from typing import List, Dict, Any


def extract_questions(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    """
    Extract questions using Python heuristics (synchronous).
    Used for initial question count. The async LLM version is used during processing.
    """
    file_type = file_type.lower().lstrip(".")
    try:
        if file_type == "docx":
            return _extract_from_docx(file_path)
        elif file_type == "xlsx":
            return _extract_from_xlsx(file_path)
        elif file_type == "txt":
            return _extract_from_txt(file_path)
        elif file_type == "pdf":
            return _extract_from_pdf(file_path)
        elif file_type == "json":
            return _extract_from_json(file_path)
        else:
            return []
    except Exception:
        return []


async def extract_questions_smart(file_path: str, file_type: str, settings_dict: dict = None) -> List[Dict[str, Any]]:
    """
    Extract questions using LLM for best quality.
    For XLSX: LLM detects layout (columns), then heuristic extracts with correct columns.
    For other formats: LLM extracts questions, matched to heuristic locations.
    Falls back to Python heuristics if LLM is unavailable.
    """
    from llm_skills import extract_questions_with_llm

    llm_settings = None
    if settings_dict:
        llm_settings = {
            "provider": settings_dict.get("llm_provider", "claude_code"),
            "ollama_url": settings_dict.get("ollama_url", "http://localhost:11434"),
            "ollama_model": settings_dict.get("llm_model", "llama3.2"),
        }

    file_type_lower = file_type.lower().lstrip(".")

    # For XLSX: use LLM to detect layout, then extract with correct columns
    if file_type_lower == "xlsx":
        result = await _extract_from_xlsx_smart(file_path, llm_settings)
        if result:
            print(f"LLM-assisted XLSX extraction: {len(result)} questions from {file_path}")
            return result
        # Fallback to heuristic
        return extract_questions(file_path, file_type)

    # For other formats: use generic LLM extraction + heuristic location matching
    doc_text = _get_document_text(file_path, file_type)
    if not doc_text:
        return extract_questions(file_path, file_type)

    llm_questions = await extract_questions_with_llm(doc_text, settings_dict=llm_settings)

    if llm_questions and len(llm_questions) > 0:
        heuristic_questions = extract_questions(file_path, file_type)
        heuristic_used = [False] * len(heuristic_questions)

        result = []
        for i, q in enumerate(llm_questions):
            q_text = q.get("question", "")
            if not q_text:
                continue

            matched = False
            for j, hq in enumerate(heuristic_questions):
                if not heuristic_used[j] and _text_matches(hq["question_text"], q_text):
                    heuristic_used[j] = True
                    result.append({
                        "question_text": q_text,
                        "section_context": hq.get("section_context", ""),
                        "location_info": hq["location_info"],
                    })
                    matched = True
                    break

            if not matched:
                location, section_ctx = _find_location_in_document(file_path, file_type, q_text, i)
                result.append({
                    "question_text": q_text,
                    "section_context": section_ctx,
                    "location_info": location,
                })

        if result:
            print(f"LLM extracted {len(result)} questions from {file_path}")
            return result

    print(f"LLM extraction failed, using heuristics for {file_path}")
    return extract_questions(file_path, file_type)


async def _extract_from_xlsx_smart(file_path: str, llm_settings: dict = None) -> List[Dict[str, Any]]:
    """LLM-assisted XLSX extraction: LLM detects table regions (possibly multiple per sheet),
    then extracts questions from the correct columns with proper answer column mapping."""
    from openpyxl import load_workbook
    from llm_skills import detect_xlsx_layout_with_llm

    try:
        wb = load_workbook(file_path)
    except Exception:
        return []

    questions = []

    for ws in wb.worksheets:
        # Build a preview — send up to 50 rows so LLM can detect multiple tables
        max_preview = ws.max_row or 50
        max_preview = min(max_preview, 50)
        preview_lines = []
        for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=max_preview, values_only=False), start=1):
            cells = []
            for cell in row:
                try:
                    col_letter = cell.column_letter
                except AttributeError:
                    continue
                val = str(cell.value).strip()[:60] if cell.value else ""
                if val:
                    cells.append(f"{col_letter}: {val}")
            if cells:
                preview_lines.append(f"Row {row_idx}: {' | '.join(cells)}")

        if not preview_lines:
            continue

        preview_text = "\n".join(preview_lines)

        # Ask LLM to detect table regions (returns a LIST)
        table_regions = None
        if llm_settings:
            table_regions = await detect_xlsx_layout_with_llm(preview_text, ws.title, llm_settings)

        if table_regions:
            for region in table_regions:
                q_col_letter = region["question_col"]
                answer_col = region["answer_col"]
                start_row = region["start_row"]
                end_row = region["end_row"]
                section_name = region.get("section", ws.title)
                q_col_idx = ord(q_col_letter.upper()) - ord('A')
                print(f"  LLM table region '{ws.title}': Q={q_col_letter}, Ans={answer_col}, rows {start_row}-{end_row}, section='{section_name}'")

                questions.extend(
                    _extract_from_region(ws, q_col_idx, answer_col, start_row, end_row, section_name, use_llm_mode=True)
                )
        else:
            # Fallback to heuristic detection (single table per sheet)
            q_col_idx, answer_col = _detect_xlsx_columns(ws)
            print(f"  Heuristic layout for '{ws.title}': questions=col{q_col_idx}, answers={answer_col}")
            questions.extend(
                _extract_from_region(ws, q_col_idx, answer_col, 2, 99999, ws.title, use_llm_mode=False)
            )

    return questions


def _extract_from_region(ws, q_col_idx: int, answer_col: str,
                         start_row: int, end_row: int, section_name: str,
                         use_llm_mode: bool = True) -> List[Dict[str, Any]]:
    """Extract questions from a specific region of a worksheet."""
    questions = []
    current_section = section_name

    for row_idx, row in enumerate(ws.iter_rows(min_row=1, values_only=False), start=1):
        if row_idx < start_row or row_idx > end_row:
            continue
        if q_col_idx >= len(row):
            continue

        cell = row[q_col_idx]
        if not cell.value:
            continue
        text = str(cell.value).strip()
        if not text or len(text) < 10:
            continue

        # Skip pure numbers
        try:
            float(text)
            continue
        except ValueError:
            pass

        # Skip header labels
        skip_lower = text.lower()
        if skip_lower in ("question", "response", "comments", "answer", "remarks",
                          "no.", "no", "#", "s.no", "sr.no", "sl.no",
                          "requirement", "requirements", "criteria", "item",
                          "section", "category", "topic", "area"):
            continue

        # Section heading detection
        if use_llm_mode:
            is_bold = cell.font and cell.font.bold
            if is_bold and len(text) < 40 and not _is_question(text):
                current_section = text
                continue
        else:
            is_q = _is_question(text)
            if not is_q:
                is_bold = cell.font and cell.font.bold
                is_instructional = skip_lower.startswith(("if ", "please ", "note:", "note ", "for ", "where ", "when "))
                looks_like_heading = (
                    len(text) < 60
                    and not text.endswith((".", ",", "?", ":"))
                    and not is_instructional
                )
                if is_bold and not is_instructional:
                    current_section = text
                elif looks_like_heading:
                    current_section = text
                continue

        questions.append({
            "question_text": text,
            "section_context": current_section,
            "location_info": {
                "type": "xlsx",
                "sheet": ws.title,
                "row": row_idx,
                "answer_col": answer_col,
            },
        })

    return questions


def _get_document_text(file_path: str, file_type: str) -> str:
    """Get plain text from document for LLM processing."""
    file_type = file_type.lower().lstrip(".")
    try:
        if file_type == "docx":
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        elif file_type == "xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True)
            lines = []
            for ws in wb:
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c) for c in row if c is not None]
                    if vals:
                        lines.append(" | ".join(vals))
            wb.close()
            return "\n".join(lines)
        elif file_type == "pdf":
            import fitz
            doc = fitz.open(file_path)
            parts = []
            # Include form field labels if present
            for page in doc:
                for widget in (page.widgets() or []):
                    label = (widget.field_label or "").strip()
                    name = (widget.field_name or "").strip()
                    if label or name:
                        parts.append(f"[Form Field] {label or name}")
            # Include page text
            for page in doc:
                t = page.get_text().strip()
                if t:
                    parts.append(t)
            doc.close()
            return "\n".join(parts)
        elif file_type == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif file_type == "json":
            import json
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
            if isinstance(data, list):
                return "\n".join(
                    str(item.get("question", item.get("q", item))) if isinstance(item, dict) else str(item)
                    for item in data
                )
            return str(data)
    except Exception:
        pass
    return ""


def _find_location_in_document(file_path: str, file_type: str, question_text: str, index: int) -> tuple:
    """Find the best matching location and section context for a question.
    Returns (location_dict, section_context_str)."""
    file_type = file_type.lower().lstrip(".")

    if file_type == "docx":
        from docx import Document
        try:
            doc = Document(file_path)
            # Build section context from heading paragraphs
            current_section = ""
            para_sections = {}
            for p_idx, para in enumerate(doc.paragraphs):
                if para.text.strip() and para.style and para.style.name and "Heading" in para.style.name:
                    current_section = para.text.strip()
                para_sections[p_idx] = current_section

            # Search in tables first
            for t_idx, table in enumerate(doc.tables):
                first_row_text = table.rows[0].cells[0].text.strip() if table.rows else ""
                for r_idx, row in enumerate(table.rows):
                    if len(row.cells) >= 2:
                        cell_text = row.cells[0].text.strip()
                        if cell_text and _text_matches(cell_text, question_text):
                            return ({"type": "docx_table", "table_index": t_idx, "row_index": r_idx, "answer_col": 1},
                                    first_row_text if not _is_question(first_row_text) else "")
            # Search in paragraphs
            for p_idx, para in enumerate(doc.paragraphs):
                if para.text.strip() and _text_matches(para.text.strip(), question_text):
                    return ({"type": "docx_paragraph", "paragraph_index": p_idx},
                            para_sections.get(p_idx, ""))
        except Exception:
            pass
        return ({"type": "docx_paragraph", "paragraph_index": index}, "")

    elif file_type == "xlsx":
        from openpyxl import load_workbook
        fallback_answer_col = "B"
        try:
            wb = load_workbook(file_path, read_only=True)
            for ws in wb.worksheets:
                # Auto-detect question and answer columns for this sheet
                q_col_idx, answer_col = _detect_xlsx_columns(ws)
                if fallback_answer_col == "B":
                    fallback_answer_col = answer_col  # Use first sheet's detection as fallback
                current_section = ws.title
                for row_idx, row in enumerate(ws.iter_rows(min_row=1, values_only=False), start=1):
                    if q_col_idx >= len(row):
                        continue
                    cell = row[q_col_idx]
                    if not cell.value:
                        continue
                    text = str(cell.value).strip()
                    if not text:
                        continue
                    if _text_matches(text, question_text):
                        wb.close()
                        return ({"type": "xlsx", "sheet": ws.title, "row": row_idx, "answer_col": answer_col},
                                current_section)
                    if not _is_question(text) and len(text) < 120:
                        current_section = text
            wb.close()
        except Exception:
            pass
        return ({"type": "xlsx", "sheet_name": "Sheet1", "row": index + 2, "col": 1, "answer_col": fallback_answer_col}, "")

    elif file_type == "json":
        return ({"type": "json", "index": index, "key": "question"}, "")

    elif file_type == "txt":
        return ({"type": "txt", "line_number": index}, "")

    elif file_type == "pdf":
        import fitz
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1):
                for widget in (page.widgets() or []):
                    label = (widget.field_label or "").strip()
                    name = (widget.field_name or "").strip()
                    if (label and _text_matches(label, question_text)) or \
                       (name and _text_matches(name, question_text)):
                        doc.close()
                        return ({"type": "pdf_form_field", "page_number": page_num, "field_name": name}, "")
            doc.close()
        except Exception:
            pass
        return ({"type": "pdf", "page_number": 1, "line_index": index}, "")

    return ({"type": file_type, "index": index}, "")


def _text_matches(doc_text: str, question_text: str) -> bool:
    """Check if document text matches a question (fuzzy: first 50 chars)."""
    a = doc_text.strip()[:50].lower()
    b = question_text.strip()[:50].lower()
    return a == b or a in b or b in a


# ---------------------------------------------------------------------------
# Python heuristic extractors (fallback)
# ---------------------------------------------------------------------------

def _is_question(text: str) -> bool:
    """Heuristic: check if text looks like a compliance question or request."""
    text = text.strip()
    if not text or len(text) < 15:
        return False
    if text.endswith("?"):
        return True
    lower = text.lower()
    q_starters = (
        "what", "how", "do ", "does", "is ", "are ", "can ", "will ",
        "where", "when", "which", "who ", "why ",
    )
    if any(lower.startswith(w) for w in q_starters):
        return True
    compliance_patterns = (
        "please provide", "please detail", "please describe", "please explain",
        "please advise", "please list", "please specify", "please confirm",
        "please indicate", "please outline", "please share",
        "describe ", "explain ", "list ", "provide ", "detail ",
        "specify ", "confirm ", "indicate ", "outline ",
    )
    if any(lower.startswith(p) for p in compliance_patterns):
        return True
    contains_patterns = (
        "please provide", "please describe", "please explain", "please detail",
        "what type", "what kind", "what method", "what integration",
        "is there", "are there", "will there",
        # Compliance requirement patterns (action-oriented, not just mentions)
        "must provide", "must describe", "must explain",
    )
    if any(p in lower for p in contains_patterns):
        return True
    # Catch requirement statements with colon format: "Topic: describe/explain/provide..."
    if ":" in text and len(text) > 30:
        after_colon = text.split(":", 1)[1].strip().lower()
        action_words = ("describe", "explain", "provide", "please", "do you", "does", "is ",
                        "are ", "can ", "how ", "what ", "where", "list ", "include")
        if any(after_colon.startswith(w) for w in action_words):
            return True
    return False


def _extract_from_docx(file_path: str) -> List[Dict[str, Any]]:
    from docx import Document
    doc = Document(file_path)
    questions = []
    seen_texts = set()
    current_section = ""

    # Track section headings from paragraphs (Heading styles)
    # Build a map of paragraph index to cumulative section context
    section_at_para = {}
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and para.style and para.style.name and "Heading" in para.style.name:
            current_section = text
        section_at_para[p_idx] = current_section

    # Reset for table processing
    current_section = ""
    for t_idx, table in enumerate(doc.tables):
        # Check if there's a heading paragraph before this table
        # Use the first row as potential section context if it's not a question
        first_row_text = table.rows[0].cells[0].text.strip() if table.rows else ""
        if first_row_text and not _is_question(first_row_text):
            current_section = first_row_text

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue
            cells = row.cells
            if len(cells) >= 2:
                q_text = cells[0].text.strip()
                if q_text and _is_question(q_text) and q_text not in seen_texts:
                    seen_texts.add(q_text)
                    questions.append({
                        "question_text": q_text,
                        "section_context": current_section,
                        "location_info": {"type": "docx_table", "table_index": t_idx, "row_index": r_idx, "answer_col": 1},
                    })

    current_section = ""
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # Update section context from headings
        if para.style and para.style.name and "Heading" in para.style.name:
            current_section = text
            continue
        if _is_question(text) and text not in seen_texts:
            seen_texts.add(text)
            questions.append({
                "question_text": text,
                "section_context": current_section,
                "location_info": {"type": "docx_paragraph", "paragraph_index": p_idx},
            })

    return questions


def _detect_xlsx_columns(ws) -> tuple:
    """Auto-detect which column has questions and which is for answers.
    Returns (question_col_index, answer_col_letter)."""
    # Check header row for clues
    header_row = list(ws.iter_rows(min_row=1, max_row=1, values_only=False))
    if not header_row:
        return (0, "B")

    q_col = 0  # default: column A
    a_col = "B"  # default: column B
    header_hints_q = ("question", "requirement", "item", "description", "criteria", "query")
    header_hints_a = ("response", "answer", "reply", "comment", "result")

    q_found = False
    a_found = False
    for cell in header_row[0]:
        if cell.value:
            val = str(cell.value).strip().lower()
            col_letter = cell.column_letter
            col_idx = cell.column - 1
            if not q_found:
                for hint in header_hints_q:
                    if hint in val:
                        q_col = col_idx
                        q_found = True
                        break
            if not a_found:
                for hint in header_hints_a:
                    if hint in val:
                        a_col = col_letter
                        a_found = True
                        break

    # If header didn't help, check first few data rows to find which column has long text
    if q_col == 0:
        for row in ws.iter_rows(min_row=2, max_row=5, values_only=False):
            for cell in row:
                if cell.value and isinstance(cell.value, str) and len(cell.value) > 30 and _is_question(cell.value):
                    q_col = cell.column - 1
                    # Answer column = next column after question
                    a_col = chr(ord('A') + q_col + 1) if q_col + 1 < 26 else "B"
                    return (q_col, a_col)

    return (q_col, a_col)


def _extract_from_xlsx(file_path: str) -> List[Dict[str, Any]]:
    from openpyxl import load_workbook
    wb = load_workbook(file_path)
    questions = []

    for ws in wb.worksheets:
        current_section = ws.title

        # Auto-detect question and answer columns
        q_col_idx, answer_col = _detect_xlsx_columns(ws)

        for row_idx, row in enumerate(ws.iter_rows(min_row=1, values_only=False), start=1):
            if q_col_idx >= len(row):
                continue
            cell = row[q_col_idx]
            if not cell.value:
                continue
            text = str(cell.value).strip()
            if not text:
                continue

            # Skip header-like values
            skip_lower = text.lower()
            if skip_lower in ("question", "response", "comments", "answer", "remarks",
                              "no.", "no", "#", "s.no", "sr.no", "sl.no",
                              "requirement", "requirements", "criteria", "item"):
                continue

            # Skip pure numbers (row numbers like 1, 2, 3)
            try:
                float(text)
                continue
            except ValueError:
                pass

            # Detect section headings
            is_q = _is_question(text)
            if not is_q:
                is_bold = cell.font and cell.font.bold
                is_instructional = skip_lower.startswith((
                    "if ", "please ", "note:", "note ", "for ", "where ", "when ",
                ))
                looks_like_heading = (
                    len(text) < 60
                    and not text.endswith((".", ",", "?", ":"))
                    and not is_instructional
                )
                if is_bold and not is_instructional:
                    current_section = text
                elif looks_like_heading:
                    current_section = text
                continue

            if is_q and row_idx >= 2:
                questions.append({
                    "question_text": text,
                    "section_context": current_section,
                    "location_info": {"type": "xlsx", "sheet": ws.title, "row": row_idx, "answer_col": answer_col},
                })
    return questions


def _extract_from_txt(file_path: str) -> List[Dict[str, Any]]:
    questions = []
    current_section = ""
    with open(file_path, "r", encoding="utf-8") as f:
        for line_idx, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue
            if not _is_question(line) and len(line) < 100:
                current_section = line
                continue
            if _is_question(line):
                questions.append({
                    "question_text": line,
                    "section_context": current_section,
                    "location_info": {"type": "txt", "line_number": line_idx},
                })
    return questions


def _extract_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    import fitz
    doc = fitz.open(file_path)
    questions = []

    # 1. Try form fields first (fillable PDFs)
    for page_num, page in enumerate(doc, start=1):
        for widget in (page.widgets() or []):
            if widget.field_type in (fitz.PDF_WIDGET_TYPE_TEXT, fitz.PDF_WIDGET_TYPE_COMBOBOX, fitz.PDF_WIDGET_TYPE_LISTBOX):
                # Use field label from tooltip, or fall back to field_name
                label = (widget.field_label or "").strip()
                tooltip = ""
                try:
                    tooltip = (widget.field_type_string or "").strip()
                except Exception:
                    pass
                field_name = (widget.field_name or "").strip()
                q_text = label or tooltip or field_name
                if q_text and len(q_text) >= 5:
                    questions.append({
                        "question_text": q_text,
                        "location_info": {
                            "type": "pdf_form_field",
                            "page_number": page_num,
                            "field_name": field_name,
                        },
                    })

    if questions:
        doc.close()
        return questions

    # 2. Fall back to text-based extraction (non-fillable PDFs)
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        for line in text.split("\n"):
            line = line.strip()
            cleaned = re.sub(r"^\d+\.\s*", "", line)
            if cleaned and _is_question(cleaned):
                questions.append({
                    "question_text": cleaned,
                    "location_info": {"type": "pdf", "page_number": page_num},
                })
    doc.close()
    return questions


# ---------------------------------------------------------------------------
# JSON extractor with fuzzy key matching
# ---------------------------------------------------------------------------

def _find_question_key(d: dict) -> tuple:
    def normalize(s):
        return s.lower().replace("_", "").replace("-", "").replace(" ", "")
    question_patterns = ["question", "questiontext", "query", "q", "ask", "prompt", "inquiry"]
    for key, val in d.items():
        if isinstance(val, str) and val.strip():
            nk = normalize(key)
            for pattern in question_patterns:
                if pattern in nk or nk in pattern:
                    return key, val.strip()
    for key, val in d.items():
        if isinstance(val, str) and val.strip().endswith("?"):
            return key, val.strip()
    for key, val in d.items():
        if isinstance(val, str) and _is_question(val):
            return key, val.strip()
    return None, None


def _find_list_key(d: dict) -> list:
    def normalize(s):
        return s.lower().replace("_", "").replace("-", "").replace(" ", "")
    list_patterns = ["question", "item", "data", "qna", "queries", "list", "entries", "rows"]
    for key, val in d.items():
        if isinstance(val, list) and len(val) > 0:
            nk = normalize(key)
            for pattern in list_patterns:
                if pattern in nk:
                    return val
    for key, val in d.items():
        if isinstance(val, list) and len(val) > 0:
            return val
    return []


def _extract_from_json(file_path: str) -> List[Dict[str, Any]]:
    import json
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        data = json.load(f)

    questions = []

    def _extract_from_list(items: list):
        for idx, item in enumerate(items):
            if isinstance(item, dict):
                found_key, q = _find_question_key(item)
                if q:
                    questions.append({
                        "question_text": q,
                        "location_info": {"type": "json", "index": idx, "key": found_key},
                    })
            elif isinstance(item, str) and _is_question(item):
                questions.append({
                    "question_text": item.strip(),
                    "location_info": {"type": "json", "index": idx},
                })

    if isinstance(data, list):
        _extract_from_list(data)
    elif isinstance(data, dict):
        q_list = _find_list_key(data)
        if q_list:
            _extract_from_list(q_list)
        else:
            found_key, q = _find_question_key(data)
            if q:
                questions.append({
                    "question_text": q,
                    "location_info": {"type": "json", "index": 0, "key": found_key},
                })

    return questions
