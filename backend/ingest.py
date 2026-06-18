"""
Compliance Pilot — Document ingestion pipeline.
All formats → Markdown → Smart chunking → Embedding → Qdrant.
"""

import os
import re
import uuid
from datetime import datetime
from typing import List, Dict, Any

import httpx
from qdrant_client import QdrantClient
from qdrant_client.models import (
    PointStruct, VectorParams, Distance,
    SparseVectorParams, SparseVector,
)
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from models import KBDocument

# Max characters per chunk — must fit within embedding model context (mxbai-embed-large: ~512 tokens)
# 1000 chars + overlap (150) + context prefix (~80) stays safely within 512 token limit
MAX_CHUNK_CHARS = 1000
CHUNK_OVERLAP_CHARS = 150


# ===========================================================================
# Step 1: Parse any format → raw text
# ===========================================================================

def _parse_to_text(file_path: str) -> str:
    """Parse any supported file format into raw text."""
    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "json":
        return _parse_json(file_path)
    elif ext == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == "md":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext in ("html", "htm"):
        return _parse_html(file_path)
    elif ext == "csv":
        return _parse_csv(file_path)

    # Try Docling first for PDF/DOCX/XLSX (outputs markdown natively)
    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(file_path)
        return result.document.export_to_markdown()
    except Exception:
        pass

    # Fallback parsers
    if ext == "docx":
        return _parse_docx(file_path)
    elif ext == "pdf":
        return _parse_pdf(file_path)
    elif ext == "xlsx":
        return _parse_xlsx(file_path)

    return ""


def _parse_json(file_path: str) -> str:
    """Parse JSON KB article into text using fuzzy key matching."""
    import json as _json
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        data = _json.load(f)

    def _normalize(s):
        return s.lower().replace("_", "").replace("-", "").replace(" ", "")

    def _find_value(d, patterns):
        for key, val in d.items():
            if isinstance(val, str) and val.strip():
                nk = _normalize(key)
                for p in patterns:
                    if p in nk or nk in p:
                        return val
        return ""

    def _extract_from_dict(d):
        parts = []

        # Title
        title = _find_value(d, ["title", "name", "heading", "subject"])
        if title:
            parts.append(f"# {title}")

        # Metadata block — URL, author, category, tags, dates
        meta_lines = []
        url = _find_value(d, ["weburl", "url", "link", "permalink", "sourceurl"])
        if url:
            # If it's a permalink slug (not a full URL), optionally prefix it with
            # a configured base. Set KB_ARTICLE_URL_BASE for your KB's article URLs.
            if not url.startswith("http"):
                base = os.getenv("KB_ARTICLE_URL_BASE", "").rstrip("/")
                if base:
                    url = f"{base}/{url.lstrip('/')}"
            meta_lines.append(f"URL: {url}")

        # Author — may be a string or a dict with 'name'
        for key, val in d.items():
            nk = _normalize(key)
            if "author" in nk:
                if isinstance(val, dict):
                    author_name = val.get("name", val.get("displayName", ""))
                    if author_name:
                        meta_lines.append(f"Author: {author_name}")
                elif isinstance(val, str) and val.strip():
                    meta_lines.append(f"Author: {val.strip()}")
                break

        # Category
        for key, val in d.items():
            nk = _normalize(key)
            if "category" in nk and nk != "categoryid" and nk != "rootcategoryid":
                if isinstance(val, dict):
                    cat_name = val.get("name", "")
                    if cat_name:
                        meta_lines.append(f"Category: {cat_name}")
                elif isinstance(val, str) and val.strip():
                    meta_lines.append(f"Category: {val.strip()}")
                break

        # Tags
        for key, val in d.items():
            nk = _normalize(key)
            if "tag" in nk:
                if isinstance(val, list) and val:
                    meta_lines.append(f"Tags: {', '.join(str(t) for t in val)}")
                elif isinstance(val, str) and val.strip():
                    meta_lines.append(f"Tags: {val.strip()}")
                break

        # Dates
        created = _find_value(d, ["createdtime", "createdat", "created", "publishdate"])
        if created:
            meta_lines.append(f"Created: {created}")
        modified = _find_value(d, ["modifiedtime", "updatedat", "modified", "lastmodified"])
        if modified:
            meta_lines.append(f"Modified: {modified}")

        if meta_lines:
            parts.append("\n".join(meta_lines))

        # Summary
        summary = _find_value(d, ["summary", "abstract", "description", "synopsis"])
        if summary and summary != title:
            parts.append(summary)

        # Main content
        content = _find_value(d, ["content", "body", "text", "answer", "detail", "article"])
        if content and content != title and content != summary:
            parts.append(content)

        # Fallback: include any string fields not already captured
        if len(parts) <= 1:
            for k, v in d.items():
                if isinstance(v, str) and len(v) > 10 and v not in parts:
                    parts.append(f"{k}: {v}")
                elif isinstance(v, dict):
                    # Recurse into nested dicts
                    for nk, nv in v.items():
                        if isinstance(nv, str) and len(nv) > 10:
                            parts.append(f"{k}.{nk}: {nv}")
                elif isinstance(v, list) and v:
                    # Capture list items
                    list_strs = [str(i) for i in v if i]
                    if list_strs:
                        parts.append(f"{k}: {', '.join(list_strs)}")

        return "\n\n".join(parts)

    if isinstance(data, dict):
        return _extract_from_dict(data)
    elif isinstance(data, list):
        return "\n\n---\n\n".join(
            _extract_from_dict(item) if isinstance(item, dict) else str(item)
            for item in data if item
        )
    return str(data)


def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            # Convert heading styles to markdown
            if p.style and p.style.name and "Heading" in p.style.name:
                level = p.style.name.replace("Heading ", "").strip()
                try:
                    hashes = "#" * int(level)
                except ValueError:
                    hashes = "#"
                parts.append(f"{hashes} {text}")
            else:
                parts.append(text)
    return "\n\n".join(parts)


def _parse_pdf(file_path: str) -> str:
    import fitz
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append(f"## Page {i + 1}\n\n{text}")
    doc.close()
    return "\n\n".join(pages)


def _parse_xlsx(file_path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"## {sheet}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                parts.append(" | ".join(vals))
    wb.close()
    return "\n\n".join(parts)


def _parse_html(file_path: str) -> str:
    """Parse HTML file into clean text, preserving structure."""
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts = []
            self._current_tag = ""
            self._skip_tags = {"script", "style", "noscript", "nav", "footer", "header"}
            self._skip = False

        def handle_starttag(self, tag, attrs):
            self._current_tag = tag
            if tag in self._skip_tags:
                self._skip = True
            elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(tag[1])
                self.parts.append("\n" + "#" * level + " ")
            elif tag in ("p", "div", "br", "li", "tr"):
                self.parts.append("\n")
            elif tag == "td" or tag == "th":
                self.parts.append(" | ")
            elif tag == "a":
                for name, val in attrs:
                    if name == "href" and val and val.startswith("http"):
                        self.parts.append(f" [{val}] ")

        def handle_endtag(self, tag):
            if tag in self._skip_tags:
                self._skip = False
            if tag in ("p", "div", "tr", "ul", "ol", "table"):
                self.parts.append("\n")
            self._current_tag = ""

        def handle_data(self, data):
            if not self._skip:
                text = data.strip()
                if text:
                    self.parts.append(text)

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html_content = f.read()

    # Extract title from <title> tag
    title = ""
    import re as _re
    title_match = _re.search(r"<title[^>]*>(.*?)</title>", html_content, _re.IGNORECASE | _re.DOTALL)
    if title_match:
        title = title_match.group(1).strip()

    parser = _TextExtractor()
    parser.feed(html_content)
    text = "".join(parser.parts).strip()

    # Add title as heading if found
    if title and not text.startswith(f"# {title}"):
        text = f"# {title}\n\n{text}"

    return text


def _parse_csv(file_path: str) -> str:
    """Parse CSV file into text with headers as context."""
    import csv as _csv
    parts = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = _csv.reader(f)
        headers = None
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
                parts.append("## Columns: " + " | ".join(row))
            else:
                if headers:
                    # Format as "header: value" pairs for better context
                    pairs = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
                    parts.append(" | ".join(pairs) if pairs else " | ".join(row))
                else:
                    parts.append(" | ".join(row))
    return "\n\n".join(parts)


# ===========================================================================
# Step 2: Convert raw text → clean Markdown
# ===========================================================================

def _convert_to_markdown(text: str, source_file: str) -> str:
    """Clean and normalize text into well-structured markdown."""
    if not text or not text.strip():
        return ""

    lines = text.strip().split("\n")
    cleaned = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append("")
            continue

        # Clean HTML tags if any
        stripped = re.sub(r"<[^>]+>", "", stripped)
        # Clean excessive whitespace
        stripped = re.sub(r"\s{3,}", "  ", stripped)

        cleaned.append(stripped)

    md = "\n".join(cleaned)

    # If no markdown headings exist, add a title from filename
    if not re.search(r"^#\s", md, re.MULTILINE):
        title = os.path.splitext(source_file)[0].replace("-", " ").replace("_", " ").title()
        md = f"# {title}\n\n{md}"

    return md.strip()


# ===========================================================================
# Step 3: Smart chunking by markdown sections (no data loss)
# ===========================================================================

def chunk_document(content: str, source_file: str,
                   max_chars: int = MAX_CHUNK_CHARS,
                   overlap_chars: int = CHUNK_OVERLAP_CHARS) -> List[Dict[str, Any]]:
    """
    Split markdown content into chunks by section boundaries.
    No data is lost — long sections are split further by character limit.
    """
    if not content or not content.strip():
        return []

    # Split by markdown headings and horizontal rules
    sections = re.split(r"\n(?=#{1,4}\s)|(?:\n---\n)", content)

    chunks = []
    chunk_index = 0

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Extract section title if present
        title_match = re.match(r"^(#{1,4})\s+(.+)", section)
        section_title = title_match.group(2).strip() if title_match else ""

        if len(section) <= max_chars:
            # Section fits in one chunk
            chunks.append({
                "text": section,
                "source_file": source_file,
                "page_number": 1,
                "section_title": section_title,
                "chunk_index": chunk_index,
                "ingested_at": datetime.utcnow().isoformat(),
            })
            chunk_index += 1
        else:
            # Section too long — split by paragraphs first, then by character limit
            sub_chunks = _split_long_section(section, section_title, max_chars)
            for sc in sub_chunks:
                sc["source_file"] = source_file
                sc["page_number"] = 1
                sc["chunk_index"] = chunk_index
                sc["ingested_at"] = datetime.utcnow().isoformat()
                chunks.append(sc)
                chunk_index += 1

    return _apply_overlap(chunks, overlap_chars=overlap_chars)


def _apply_overlap(chunks: List[Dict[str, Any]], overlap_chars: int = CHUNK_OVERLAP_CHARS) -> List[Dict[str, Any]]:
    """Add overlap text from previous chunk's tail to current chunk's beginning."""
    if overlap_chars <= 0 or len(chunks) < 2:
        return chunks
    for i in range(1, len(chunks)):
        prev_text = chunks[i - 1]["text"]
        if len(prev_text) > overlap_chars:
            overlap_text = prev_text[-overlap_chars:]
            # Find a clean word boundary
            space_idx = overlap_text.find(' ')
            if space_idx > 0:
                overlap_text = overlap_text[space_idx + 1:]
            chunks[i]["text"] = f"...{overlap_text}\n\n{chunks[i]['text']}"
    return chunks


def _split_long_section(text: str, section_title: str,
                        max_chars: int) -> List[Dict]:
    """Split a long section into smaller chunks without losing any data.
    Tries paragraph boundaries first, falls back to sentence/word boundaries."""
    paragraphs = text.split("\n\n")
    sub_chunks = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # If adding this paragraph stays under limit, append it
        if len(current) + len(para) + 2 <= max_chars:
            current = f"{current}\n\n{para}" if current else para
        else:
            # Save current chunk if non-empty
            if current.strip():
                sub_chunks.append({
                    "text": current.strip(),
                    "section_title": section_title,
                })

            # If the single paragraph itself is too long, split by sentences
            if len(para) > max_chars:
                sentence_chunks = _split_by_sentences(para, max_chars)
                for sc in sentence_chunks:
                    sub_chunks.append({
                        "text": sc,
                        "section_title": section_title,
                    })
                current = ""
            else:
                current = para

    # Don't forget the last chunk
    if current.strip():
        sub_chunks.append({
            "text": current.strip(),
            "section_title": section_title,
        })

    return sub_chunks


def _split_by_sentences(text: str, max_chars: int) -> List[str]:
    """Split text by sentence boundaries. Last resort: split by words."""
    # Try splitting by sentences (period + space)
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks = []
    current = ""

    for sent in sentences:
        if len(current) + len(sent) + 1 <= max_chars:
            current = f"{current} {sent}" if current else sent
        else:
            if current.strip():
                chunks.append(current.strip())
            # If single sentence is still too long, hard split by words
            if len(sent) > max_chars:
                words = sent.split()
                word_chunk = ""
                for w in words:
                    if len(word_chunk) + len(w) + 1 <= max_chars:
                        word_chunk = f"{word_chunk} {w}" if word_chunk else w
                    else:
                        if word_chunk.strip():
                            chunks.append(word_chunk.strip())
                        word_chunk = w
                if word_chunk.strip():
                    chunks.append(word_chunk.strip())
                current = ""
            else:
                current = sent

    if current.strip():
        chunks.append(current.strip())

    return chunks


# ===========================================================================
# Step 4: Embeddings
# ===========================================================================

def _no_proxy_client(timeout: float = 60.0) -> httpx.AsyncClient:
    """Create an httpx client that bypasses corporate proxy for local services."""
    transport = httpx.AsyncHTTPTransport()
    return httpx.AsyncClient(
        timeout=timeout,
        mounts={"http://localhost": transport, "http://127.0.0.1": transport},
    )


async def get_embedding(text: str, ollama_url: str, model: str) -> List[float]:
    """Get embedding vector from Ollama with retry on failure."""
    import asyncio
    for attempt in range(3):
        try:
            async with _no_proxy_client(60.0) as client:
                resp = await client.post(
                    f"{ollama_url}/api/embeddings",
                    json={"model": model, "prompt": text},
                )
                resp.raise_for_status()
                return resp.json()["embedding"]
        except Exception:
            if attempt < 2:
                await asyncio.sleep(1 * (attempt + 1))  # 1s, 2s backoff
            else:
                raise


# Cache the sparse model instance to avoid reloading on every call
_sparse_model = None


def _get_sparse_embedding_sync(text: str) -> dict:
    """Get sparse embedding from fastembed BM42 model (CPU-bound, synchronous)."""
    global _sparse_model
    if _sparse_model is None:
        from fastembed import SparseTextEmbedding
        _sparse_model = SparseTextEmbedding(model_name="Qdrant/bm42-all-minilm-l6-v2-attentions")
    results = list(_sparse_model.embed([text]))
    if results:
        sparse = results[0]
        return {"indices": sparse.indices.tolist(), "values": sparse.values.tolist()}
    return {"indices": [], "values": []}


async def get_sparse_embedding(text: str) -> dict:
    """Async wrapper — runs CPU-bound sparse embedding in a thread to avoid blocking event loop."""
    import asyncio
    return await asyncio.to_thread(_get_sparse_embedding_sync, text)


def get_sparse_embedding_sync(text: str) -> dict:
    """Synchronous version for use inside thread-based batch processing."""
    return _get_sparse_embedding_sync(text)


# ===========================================================================
# Qdrant collection management
# ===========================================================================

def get_collection_name(version: int) -> str:
    return f"policy_v{version}"


async def detect_embedding_dims(ollama_url: str, embed_model: str) -> int:
    """Auto-detect embedding dimensions by running a test embedding."""
    try:
        test_vec = await get_embedding("test", ollama_url, embed_model)
        return len(test_vec)
    except Exception:
        return 1024  # fallback default


def get_or_create_collection(qdrant_url: str, version: int, vector_size: int = 1024):
    client = QdrantClient(url=qdrant_url)
    collection_name = get_collection_name(version)
    if not client.collection_exists(collection_name):
        client.create_collection(
            collection_name=collection_name,
            vectors_config={
                "dense": VectorParams(size=vector_size, distance=Distance.COSINE),
            },
            sparse_vectors_config={
                "sparse": SparseVectorParams(),
            },
        )
    return client, collection_name


# ===========================================================================
# Version management
# ===========================================================================

async def get_current_version(db: AsyncSession) -> int:
    """Get the active KB version from kb_versions table."""
    from models import KBVersion
    result = await db.execute(
        select(KBVersion).where(KBVersion.is_active == True).limit(1)
    )
    active = result.scalar_one_or_none()
    if active:
        return active.version
    # Fallback: settings table
    from models import Settings as SettingsModel
    result = await db.execute(
        select(SettingsModel).where(SettingsModel.key == "kb_version")
    )
    setting = result.scalar_one_or_none()
    if setting:
        return int(setting.value)
    # Last fallback: max doc version
    result = await db.execute(select(func.max(KBDocument.version)))
    version = result.scalar()
    return version or 1


async def create_new_version(db: AsyncSession, name: str = None, embed_model: str = None) -> int:
    """Create a new KB version with metadata."""
    from models import KBVersion, Settings as SettingsModel
    from datetime import datetime as _dt

    current = await get_current_version(db)
    new_ver = current + 1

    if not name:
        name = f"Version {new_ver}" + (f" ({embed_model})" if embed_model else "")

    # Create KBVersion record
    kv = KBVersion(
        version=new_ver,
        name=name,
        embed_model=embed_model or "unknown",
        doc_count=0,
        is_active=False,
    )
    db.add(kv)

    # Also update settings table for backward compat
    result = await db.execute(
        select(SettingsModel).where(SettingsModel.key == "kb_version")
    )
    setting = result.scalar_one_or_none()
    if setting:
        setting.value = str(new_ver)
        setting.updated_at = _dt.utcnow()
    else:
        db.add(SettingsModel(key="kb_version", value=str(new_ver)))

    await db.commit()
    return new_ver


# ===========================================================================
# Main pipeline: Parse → Markdown → Chunk → Embed → Store
# ===========================================================================

# ===========================================================================
# Q&A-aware ingestion for previously answered questionnaires (Option B)
# ===========================================================================


async def _extract_qa_chunks(file_path: str, filename: str,
                              llm_settings: dict,
                              max_chunk_chars: int = MAX_CHUNK_CHARS) -> List[Dict[str, Any]]:
    """Extract Q&A pairs from a filled questionnaire and format as labeled chunks."""
    ext = file_path.rsplit(".", 1)[-1].lower()

    pairs = []
    if ext == "xlsx":
        pairs = await _extract_qa_from_xlsx(file_path, llm_settings)
    elif ext == "docx":
        pairs = _extract_qa_from_docx(file_path)
    elif ext in ("pdf", "txt"):
        raw_text = _parse_to_text(file_path)
        return await _extract_qa_from_text_llm(raw_text, filename, llm_settings, max_chunk_chars)

    if not pairs:
        return []

    chunks = []
    for i, pair in enumerate(pairs):
        q = pair.get("question", "").strip()
        a = pair.get("answer", "").strip()
        if not q or not a or len(a) < 3:
            continue

        chunk_text = (
            "[APPROVED PRIOR ANSWER]\n"
            f"Question the customer asked: {q}\n"
            f"Our approved response: {a}"
        )

        # If chunk is too long, truncate answer with note
        if len(chunk_text) > max_chunk_chars * 2:
            max_a = max_chunk_chars * 2 - len(q) - 100
            a = a[:max_a] + "..."
            chunk_text = (
                "[APPROVED PRIOR ANSWER]\n"
                f"Question the customer asked: {q}\n"
                f"Our approved response: {a}"
            )

        chunks.append({
            "text": chunk_text,
            "source_file": filename,
            "page_number": pair.get("page", 1),
            "section_title": pair.get("section", ""),
            "chunk_index": i,
            "ingested_at": datetime.utcnow().isoformat(),
            "source_type": "prior_qa",
            "original_question": q,
        })

    return chunks


async def _extract_qa_from_xlsx(file_path: str, llm_settings: dict) -> List[Dict]:
    """Extract Q&A pairs from an XLSX using LLM layout detection."""
    from openpyxl import load_workbook
    from llm_skills import detect_xlsx_layout_with_llm

    wb = load_workbook(file_path, data_only=True)
    pairs = []

    for ws in wb.worksheets:
        max_preview = min(ws.max_row or 50, 50)
        preview_lines = []
        for row_idx, row in enumerate(
            ws.iter_rows(min_row=1, max_row=max_preview, values_only=False), start=1
        ):
            cells = []
            for cell in row:
                col_letter = cell.column_letter if hasattr(cell, "column_letter") else ""
                if not col_letter:
                    continue
                val = str(cell.value).strip()[:60] if cell.value else ""
                if val:
                    cells.append(f"{col_letter}: {val}")
            if cells:
                preview_lines.append(f"Row {row_idx}: {' | '.join(cells)}")

        if not preview_lines:
            continue

        table_regions = await detect_xlsx_layout_with_llm(
            "\n".join(preview_lines), ws.title, llm_settings
        )

        if not table_regions:
            continue

        for region in table_regions:
            q_col_idx = ord(region["question_col"].upper()) - ord("A")
            a_col_idx = ord(region["answer_col"].upper()) - ord("A")
            start_row = region["start_row"]
            end_row = region["end_row"]
            section = region.get("section", ws.title)

            for row_idx, row in enumerate(ws.iter_rows(min_row=1, values_only=True), start=1):
                if row_idx < start_row or row_idx > end_row:
                    continue
                if q_col_idx >= len(row) or a_col_idx >= len(row):
                    continue
                q_val = str(row[q_col_idx] or "").strip()
                a_val = str(row[a_col_idx] or "").strip()
                if q_val and a_val and len(q_val) > 10 and len(a_val) > 3:
                    pairs.append({
                        "question": q_val,
                        "answer": a_val,
                        "section": section,
                        "page": 1,
                    })

    wb.close()
    return pairs


def _extract_qa_from_docx(file_path: str) -> List[Dict]:
    """Extract Q&A pairs from a DOCX questionnaire (tables + paragraph patterns)."""
    from docx import Document

    doc = Document(file_path)
    pairs = []
    current_section = ""

    # 1. Extract from tables
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]

        # Detect Q and A columns
        q_col, a_col = 0, 1
        for i, h in enumerate(header_cells):
            if any(w in h for w in ("question", "requirement", "item", "description", "query")):
                q_col = i
            if any(w in h for w in ("response", "answer", "reply", "comment", "vendor")):
                a_col = i

        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) > max(q_col, a_col):
                q_text = cells[q_col].text.strip()
                a_text = cells[a_col].text.strip()
                if q_text and a_text and len(q_text) > 10 and len(a_text) > 3:
                    pairs.append({
                        "question": q_text,
                        "answer": a_text,
                        "section": current_section,
                        "page": 1,
                    })

    # 2. Extract from paragraph patterns: "Q: ... A: ..."
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p.text.strip()

        if p.style and p.style.name and "Heading" in p.style.name:
            current_section = text
            i += 1
            continue

        if text.lower().startswith(("q:", "question:")) and i + 1 < len(paragraphs):
            next_text = paragraphs[i + 1].text.strip()
            if next_text.lower().startswith(("a:", "answer:", "response:")):
                q = re.sub(r"^(?:q|question)\s*:\s*", "", text, flags=re.IGNORECASE)
                a = re.sub(r"^(?:a|answer|response)\s*:\s*", "", next_text, flags=re.IGNORECASE)
                if q and a:
                    pairs.append({"question": q, "answer": a, "section": current_section, "page": 1})
                i += 2
                continue
        i += 1

    return pairs


async def _extract_qa_from_text_llm(raw_text: str, filename: str,
                                     llm_settings: dict,
                                     max_chunk_chars: int = MAX_CHUNK_CHARS) -> List[Dict[str, Any]]:
    """Use LLM to extract Q&A pairs from unstructured text (PDF/TXT fallback)."""
    import json as _json
    from llm_skills import _call_llm

    max_chars = 12000
    truncated = raw_text[:max_chars] if len(raw_text) > max_chars else raw_text

    prompt = (
        "This document is a previously answered compliance questionnaire. "
        "Extract every question-answer pair where BOTH a question and an answer are present.\n\n"
        "Rules:\n"
        "- Extract only pairs with both question and answer\n"
        "- Skip unanswered questions, headers, metadata\n"
        "- Keep full text of both question and answer\n\n"
        "Return ONLY a valid JSON array:\n"
        '[{"question": "...", "answer": "..."}]\n\n'
        f'Document:\n"""\n{truncated}\n"""\n\n'
        "Return ONLY the JSON array."
    )

    provider = llm_settings.get("provider", "claude_code")
    ollama_url = llm_settings.get("ollama_url", "http://localhost:11434")
    llm_model = llm_settings.get("ollama_model", "llama3.2")
    response = await _call_llm(prompt, provider, ollama_url, llm_model, timeout=60)
    if not response:
        return []

    try:
        from json_utils import extract_json
        raw_pairs = extract_json(response, expect=list)
        if raw_pairs:
            chunks = []
            for i, p in enumerate(raw_pairs):
                if not isinstance(p, dict) or not p.get("question") or not p.get("answer"):
                    continue
                q = p["question"].strip()
                a = p["answer"].strip()
                if len(q) < 10 or len(a) < 3:
                    continue
                chunk_text = (
                    "[APPROVED PRIOR ANSWER]\n"
                    f"Question the customer asked: {q}\n"
                    f"Our approved response: {a}"
                )
                chunks.append({
                    "text": chunk_text,
                    "source_file": filename,
                    "page_number": 1,
                    "section_title": "",
                    "chunk_index": i,
                    "ingested_at": datetime.utcnow().isoformat(),
                    "source_type": "prior_qa",
                    "original_question": q,
                })
            return chunks
    except Exception as e:
        print(f"LLM Q&A extraction failed: {e}")

    return []


async def _embed_and_store_chunks(
    chunks: List[Dict[str, Any]], filename: str, version: int,
    qdrant_url: str, ollama_url: str, embed_model: str,
    embed_concurrency: int = 1,
) -> int:
    """Embed a list of pre-built chunks and store in Qdrant. Shared by normal and Q&A paths."""
    if not chunks:
        return 0

    vector_dims = await detect_embedding_dims(ollama_url, embed_model)
    client, collection_name = get_or_create_collection(qdrant_url, version, vector_size=vector_dims)

    doc_name = os.path.splitext(filename)[0].replace("-", " ").replace("_", " ").title()

    import asyncio as _aio

    semaphore = _aio.Semaphore(max(1, embed_concurrency))
    points = []
    points_lock = _aio.Lock()

    async def _embed_chunk(chunk):
        async with semaphore:
            embed_text = f"Document: {doc_name}"
            if chunk.get("section_title"):
                embed_text += f" | Section: {chunk['section_title']}"
            embed_text += f"\n\n{chunk['text']}"

            embedding = await get_embedding(embed_text, ollama_url, embed_model)
            sparse = await get_sparse_embedding(chunk["text"])

            point = PointStruct(
                id=str(uuid.uuid4()),
                vector={
                    "dense": embedding,
                    "sparse": SparseVector(
                        indices=sparse["indices"],
                        values=sparse["values"],
                    ),
                },
                payload={
                    "text": chunk["text"],
                    "source_file": chunk["source_file"],
                    "source_url": chunk.get("source_url", ""),
                    "page_number": chunk["page_number"],
                    "section_title": chunk.get("section_title", ""),
                    "chunk_index": chunk["chunk_index"],
                    "version": version,
                    "ingested_at": chunk["ingested_at"],
                    "source_type": chunk.get("source_type", "policy"),
                    "original_question": chunk.get("original_question", ""),
                },
            )
            async with points_lock:
                points.append(point)

    if embed_concurrency <= 1:
        for chunk in chunks:
            await _embed_chunk(chunk)
    else:
        await _aio.gather(*[_embed_chunk(chunk) for chunk in chunks])

    batch_size = 100
    for i in range(0, len(points), batch_size):
        batch = points[i:i + batch_size]
        client.upsert(collection_name=collection_name, points=batch)

    return len(points)


async def ingest_document(
    file_path: str, filename: str, version: int, db: AsyncSession,
    qdrant_url: str, ollama_url: str, embed_model: str,
    llm_model: str = "llama3.2:3b",
    use_llm_markdown: bool = False,
    max_chunk_chars: int = MAX_CHUNK_CHARS,
    embed_concurrency: int = 1,
    chunk_overlap: int = CHUNK_OVERLAP_CHARS,
    is_questionnaire: bool = False,
    llm_provider: str = "claude_code",
    llm_ollama_url: str = "http://localhost:11434",
) -> int:
    """Full ingestion pipeline: Parse → Markdown → Chunk → Embed → Store."""

    # Step 1: Parse to raw text
    raw_text = _parse_to_text(file_path)
    if not raw_text.strip():
        return 0

    # Q&A path: extract structured Q&A pairs instead of normal chunking
    if is_questionnaire:
        llm_settings = {"provider": llm_provider, "ollama_url": llm_ollama_url, "ollama_model": llm_model}
        qa_chunks = await _extract_qa_chunks(file_path, filename, llm_settings, max_chunk_chars)
        if not qa_chunks:
            # Fallback: try LLM extraction from raw text
            qa_chunks = await _extract_qa_from_text_llm(raw_text, filename, llm_settings, max_chunk_chars)
        if qa_chunks:
            return await _embed_and_store_chunks(
                qa_chunks, filename, version, qdrant_url, ollama_url, embed_model, embed_concurrency
            )
        # If both fail, fall through to normal ingestion
        print(f"Q&A extraction failed for {filename}, falling back to normal ingestion")

    # Step 2: Convert to clean markdown
    # Uses Ollama for markdown conversion (free, saves Claude tokens)
    # If LLM fails for any reason, Python fallback ensures ingestion still works
    markdown = None
    if use_llm_markdown:
        try:
            from llm_skills import convert_to_markdown_with_llm
            ollama_settings = {
                "provider": "ollama",
                "ollama_url": ollama_url,
                "ollama_model": llm_model,
            }
            markdown = await convert_to_markdown_with_llm(raw_text, filename, settings_dict=ollama_settings)
            # Validate LLM output — reject if too short compared to raw (data loss)
            if markdown and len(markdown.strip()) < len(raw_text.strip()) * 0.3:
                print(f"LLM markdown too short for {filename} ({len(markdown)} vs {len(raw_text)}), using Python fallback")
                markdown = None
        except Exception as e:
            print(f"LLM markdown failed for {filename}: {e}")
            markdown = None

    # Python fallback — always works, never fails
    if not markdown or not markdown.strip():
        markdown = _convert_to_markdown(raw_text, filename)

    if not markdown.strip():
        return 0

    # Extract source URL from the markdown (if present, e.g., from JSON/TXT metadata)
    source_url = ""
    url_match = re.search(r"URL:\s*(https?://[^\s]+)", markdown)
    if url_match:
        source_url = url_match.group(1).strip()

    # Step 3: Smart chunk by sections
    chunks = chunk_document(markdown, filename, max_chars=max_chunk_chars, overlap_chars=chunk_overlap)
    if not chunks:
        return 0

    # Step 4: Embed and store — inject source_url into chunks, then use shared function
    if source_url:
        for chunk in chunks:
            chunk["source_url"] = source_url

    return await _embed_and_store_chunks(
        chunks, filename, version, qdrant_url, ollama_url, embed_model, embed_concurrency
    )
