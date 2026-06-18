"""
Shared pytest fixtures for backend tests.
All external services (Qdrant, Ollama) are mocked — tests run without Docker.
"""

import os
import sys
import json
import asyncio
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from datetime import datetime

import pytest
import pytest_asyncio
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
from sqlalchemy.orm import sessionmaker
from httpx import AsyncClient, ASGITransport

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "backend"))

from models import Base, User, AuditLog, KBDocument, ProcessingJob, Settings
from auth import hash_password


# ---------------------------------------------------------------------------
# Event loop
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def event_loop():
    loop = asyncio.new_event_loop()
    yield loop
    loop.close()


# ---------------------------------------------------------------------------
# Test database
# ---------------------------------------------------------------------------

@pytest_asyncio.fixture
async def db_engine():
    """Create a fresh in-memory SQLite database for each test."""
    engine = create_async_engine("sqlite+aiosqlite:///:memory:", echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture
async def db_session(db_engine):
    """Provide an async database session."""
    async_session_factory = sessionmaker(
        db_engine, class_=AsyncSession, expire_on_commit=False
    )
    async with async_session_factory() as session:
        yield session


@pytest_asyncio.fixture
async def db_with_admin(db_session):
    """Database session pre-loaded with default admin user."""
    admin = User(
        username="admin",
        hashed_password=hash_password("admin123"),
        role="admin",
        must_change_password=True,
    )
    db_session.add(admin)
    await db_session.commit()
    await db_session.refresh(admin)
    return db_session, admin


@pytest_asyncio.fixture
async def db_with_users(db_session):
    """Database session with admin + regular user."""
    admin = User(
        username="admin",
        hashed_password=hash_password("admin123"),
        role="admin",
        must_change_password=False,
    )
    user = User(
        username="testuser",
        hashed_password=hash_password("user123"),
        role="user",
        must_change_password=False,
    )
    db_session.add_all([admin, user])
    await db_session.commit()
    await db_session.refresh(admin)
    await db_session.refresh(user)
    return db_session, admin, user


# ---------------------------------------------------------------------------
# Mock Qdrant
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_qdrant():
    """Mock Qdrant client that returns predictable search results."""
    client = MagicMock()

    # collection_exists
    client.collection_exists.return_value = True

    # search results
    mock_point = MagicMock()
    mock_point.score = 0.88
    mock_point.payload = {
        "text": "Our data retention policy requires that all customer data be retained for a minimum of 7 years.",
        "source_file": "data_retention_policy.pdf",
        "page_number": 3,
        "section_title": "Data Retention",
        "version": 1,
        "ingested_at": "2024-01-01T00:00:00",
        "chunk_index": 0,
    }

    mock_point_low = MagicMock()
    mock_point_low.score = 0.45
    mock_point_low.payload = {
        "text": "General company overview and mission statement.",
        "source_file": "company_overview.pdf",
        "page_number": 1,
        "section_title": "Overview",
        "version": 1,
        "ingested_at": "2024-01-01T00:00:00",
        "chunk_index": 0,
    }

    mock_point_medium = MagicMock()
    mock_point_medium.score = 0.72
    mock_point_medium.payload = {
        "text": "Access control measures should be implemented for sensitive data areas.",
        "source_file": "access_control.pdf",
        "page_number": 5,
        "section_title": "Access Control",
        "version": 1,
        "ingested_at": "2024-01-01T00:00:00",
        "chunk_index": 1,
    }

    # Default: return high-confidence results
    client.search.return_value = [mock_point]
    client.query_points.return_value = MagicMock(points=[mock_point])

    # Store variants for tests to swap in
    client._mock_points = {
        "high": [mock_point],
        "low": [mock_point_low],
        "medium": [mock_point_medium],
        "empty": [],
    }

    return client


# ---------------------------------------------------------------------------
# Mock Ollama
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_ollama_response():
    """Factory for mock Ollama HTTP responses."""
    def _make(answer="Our policy states data must be retained for 7 years.", status=200):
        resp = MagicMock()
        resp.status_code = status
        resp.json.return_value = {
            "model": "llama3.2",
            "response": answer,
            "done": True,
        }
        return resp
    return _make


@pytest.fixture
def mock_embedding_response():
    """Mock embedding response from Ollama."""
    resp = MagicMock()
    resp.status_code = 200
    resp.json.return_value = {
        "embedding": [0.1] * 1024,
    }
    return resp


# ---------------------------------------------------------------------------
# Sample files fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX with a Q&A table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Compliance Questionnaire", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Question"
    table.cell(0, 1).text = "Answer"
    table.cell(1, 0).text = "What is your data retention policy?"
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "How do you handle data breaches?"
    table.cell(2, 1).text = ""
    table.cell(3, 0).text = "What encryption standards do you use?"
    table.cell(3, 1).text = ""

    path = tmp_path / "sample_questionnaire.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_xlsx(tmp_path):
    """Create a sample XLSX questionnaire."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire"
    ws["A1"] = "Question"
    ws["B1"] = "Answer"
    ws["A2"] = "What is your data retention policy?"
    ws["B2"] = ""
    ws["A3"] = "Do you conduct regular security audits?"
    ws["B3"] = ""
    ws["A4"] = "How is access control managed?"
    ws["B4"] = ""

    path = tmp_path / "sample_questionnaire.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def sample_txt(tmp_path):
    """Create a sample TXT questionnaire."""
    content = """What is your data retention policy?
How do you handle data breaches?
What encryption standards do you use?
Do you have a disaster recovery plan?
"""
    path = tmp_path / "sample_questionnaire.txt"
    path.write_text(content)
    return path


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF with questions."""
    import fitz  # pymupdf

    doc = fitz.open()
    page = doc.new_page()
    text = "Compliance Questionnaire\n\n1. What is your data retention policy?\n2. How do you handle data breaches?\n3. What encryption standards do you use?"
    page.insert_text((72, 72), text, fontsize=12)
    path = tmp_path / "sample_questionnaire.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_fillable_pdf(tmp_path):
    """Create a sample fillable PDF with form fields."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()

    # Add title text
    page.insert_text((72, 50), "Compliance Questionnaire Form", fontsize=14)

    # Add text form fields with labels
    fields = [
        ("data_retention_policy", "What is your data retention policy?", 100),
        ("breach_handling", "How do you handle data breaches?", 160),
        ("encryption_standards", "What encryption standards do you use?", 220),
    ]

    for field_name, label, y_pos in fields:
        # Add label text above the field
        page.insert_text((72, y_pos), label, fontsize=10)
        # Create text form field widget
        widget = fitz.Widget()
        widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        widget.field_name = field_name
        widget.field_label = label
        widget.rect = fitz.Rect(72, y_pos + 5, 500, y_pos + 25)
        widget.field_value = ""
        widget.text_maxlen = 500
        page.add_widget(widget)

    path = tmp_path / "sample_fillable.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_policy_docx(tmp_path):
    """Create a sample KB policy document."""
    from docx import Document

    doc = Document()
    doc.add_heading("Data Retention Policy", level=1)
    doc.add_paragraph(
        "All customer data must be retained for a minimum of 7 years "
        "in accordance with regulatory requirements. Data older than 7 years "
        "may be archived or securely deleted following the approved data "
        "destruction procedures."
    )
    doc.add_heading("Data Breach Handling", level=1)
    doc.add_paragraph(
        "In the event of a data breach, the security team must be notified "
        "within 24 hours. A full investigation must be conducted and affected "
        "parties must be notified within 72 hours as required by GDPR."
    )

    path = tmp_path / "sample_policy.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty DOCX with no questions."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This document contains no questions.")
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Settings dict fixture
# ---------------------------------------------------------------------------

@pytest.fixture
def settings_dict():
    """Default settings dictionary for testing."""
    return {
        "ollama_url": "http://localhost:11434",
        "llm_model": "llama3.2",
        "embed_model": "nomic-embed-text",
        "confidence_auto_fill": 0.82,
        "confidence_flag": 0.65,
        "max_chunks": 5,
        "hybrid_alpha": 0.5,
        "qdrant_url": "http://localhost:6333",
    }


# ---------------------------------------------------------------------------
# FastAPI test client
# ---------------------------------------------------------------------------

@pytest_asyncio.fixture
async def app_client(db_engine):
    """Create a test client with mocked database."""
    import models as models_module
    from main import app, get_db

    async_session_factory = sessionmaker(
        db_engine, class_=AsyncSession, expire_on_commit=False
    )

    async def override_get_db():
        async with async_session_factory() as session:
            yield session

    app.dependency_overrides[get_db] = override_get_db

    # Also patch models.async_session so auth.py uses the test DB
    original_async_session = models_module.async_session
    models_module.async_session = async_session_factory

    # Create admin user
    async with async_session_factory() as session:
        admin = User(
            username="admin",
            hashed_password=hash_password("admin123"),
            role="admin",
            must_change_password=False,
        )
        session.add(admin)
        await session.commit()

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        yield client

    app.dependency_overrides.clear()
    models_module.async_session = original_async_session


@pytest_asyncio.fixture
async def auth_headers(app_client):
    """Authenticate as admin. Auth is cookie-based, so the cookies are stored
    on app_client's jar; we return the CSRF header needed for mutating requests."""
    await app_client.post(
        "/api/auth/login",
        json={"username": "admin", "password": "admin123"},
    )
    csrf = app_client.cookies.get("csrf_token")
    return {"X-CSRF-Token": csrf} if csrf else {}
