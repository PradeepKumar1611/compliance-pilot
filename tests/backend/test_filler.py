"""
Tests for filler.py — fill answers into questionnaire documents.
"""

import os
import pytest
from pathlib import Path
from unittest.mock import patch
from datetime import datetime

from filler import fill_document, _format_answer, _fill_docx, _fill_xlsx, _fill_txt_as_docx, _fill_pdf


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_qa_pairs(answers=None, tiers=None):
    """Build sample qa_pairs for DOCX table filling."""
    if answers is None:
        answers = [
            "Data must be retained for 7 years.",
            "Notify security team within 24 hours.",
            "AES-256 encryption is used.",
        ]
    if tiers is None:
        tiers = ["auto_fill", "auto_fill", "auto_fill"]

    return [
        {
            "question_text": f"Question {i+1}?",
            "answer_text": ans,
            "confidence_score": 0.9,
            "confidence_tier": tier,
            "location_info": {
                "type": "docx_table",
                "table_index": 0,
                "row_index": i + 1,
                "answer_col": 1,
            },
            "sources": [{"source_file": "policy.pdf", "page_number": 1}],
        }
        for i, (ans, tier) in enumerate(zip(answers, tiers))
    ]


def _make_xlsx_qa_pairs():
    """Build sample qa_pairs for XLSX filling."""
    return [
        {
            "question_text": "What is your policy?",
            "answer_text": "7-year retention.",
            "confidence_score": 0.9,
            "confidence_tier": "auto_fill",
            "location_info": {"type": "xlsx", "sheet": "Questionnaire", "row": 2, "answer_col": "B"},
            "sources": [],
        },
        {
            "question_text": "Do you audit?",
            "answer_text": "Yes, annually.",
            "confidence_score": 0.85,
            "confidence_tier": "auto_fill",
            "location_info": {"type": "xlsx", "sheet": "Questionnaire", "row": 3, "answer_col": "B"},
            "sources": [],
        },
    ]


# ---------------------------------------------------------------------------
# Happy path
# ---------------------------------------------------------------------------


class TestFillDocxTable:
    """Test filling answers into DOCX table cells."""

    @pytest.mark.asyncio
    async def test_fill_docx_table(self, sample_docx, tmp_path):
        qa_pairs = _make_qa_pairs()

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_docx(str(sample_docx), qa_pairs, version=1)

        assert os.path.exists(output)

        from docx import Document
        doc = Document(output)
        table = doc.tables[0]

        # Question cells should be unchanged
        assert "data retention" in table.cell(1, 0).text.lower()
        # Answer cells should have content
        assert table.cell(1, 1).text.strip() != ""
        assert "7 years" in table.cell(1, 1).text


class TestFillXlsx:
    """Test filling answers into XLSX cells."""

    @pytest.mark.asyncio
    async def test_fill_xlsx(self, sample_xlsx, tmp_path):
        qa_pairs = _make_xlsx_qa_pairs()

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_xlsx(str(sample_xlsx), qa_pairs, version=1)

        assert os.path.exists(output)

        from openpyxl import load_workbook
        wb = load_workbook(output)
        ws = wb.active
        # auto_fill answers get source citation appended
        assert "7-year retention." in ws["B2"].value
        assert "Yes, annually." in ws["B3"].value


class TestFillTxtGeneratesDocx:
    """Test that TXT input generates a new DOCX output."""

    @pytest.mark.asyncio
    async def test_fill_txt_generates_docx(self, sample_txt, tmp_path):
        qa_pairs = [
            {
                "question_text": "What is your data retention policy?",
                "answer_text": "7 years.",
                "confidence_score": 0.9,
                "confidence_tier": "auto_fill",
                "location_info": {"type": "txt", "line_number": 1},
                "sources": [],
            },
        ]

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_txt_as_docx(str(sample_txt), qa_pairs, version=1)

        assert output.endswith(".docx")
        assert os.path.exists(output)

        from docx import Document
        doc = Document(output)
        assert len(doc.tables) >= 1


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


class TestPreserveFormatting:
    """Test that original DOCX formatting is preserved."""

    @pytest.mark.asyncio
    async def test_preserve_formatting(self, tmp_path):
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        heading = doc.add_heading("Compliance Questionnaire", level=1)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        hdr_cell = table.cell(0, 0)
        hdr_cell.text = ""
        p = hdr_cell.paragraphs[0]
        run = p.add_run("Question")
        run.bold = True

        table.cell(0, 1).text = "Answer"
        table.cell(1, 0).text = "What is your policy?"
        table.cell(1, 1).text = ""

        src_path = tmp_path / "formatted.docx"
        doc.save(str(src_path))

        qa_pairs = [{
            "question_text": "What is your policy?",
            "answer_text": "Our policy is XYZ.",
            "confidence_score": 0.9,
            "confidence_tier": "auto_fill",
            "location_info": {"type": "docx_table", "table_index": 0, "row_index": 1, "answer_col": 1},
            "sources": [],
        }]

        out_dir = tmp_path / "output"
        out_dir.mkdir()
        with patch("filler.OUTPUT_DIR", out_dir):
            output = _fill_docx(str(src_path), qa_pairs, version=1)

        filled_doc = Document(output)
        assert filled_doc.paragraphs[0].text == "Compliance Questionnaire"
        hdr_para = filled_doc.tables[0].cell(0, 0).paragraphs[0]
        assert any(run.bold for run in hdr_para.runs)


class TestFooterAdded:
    """Test that filled documents have Compliance Pilot footer."""

    @pytest.mark.asyncio
    async def test_footer_added(self, sample_docx, tmp_path):
        qa_pairs = _make_qa_pairs()

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_docx(str(sample_docx), qa_pairs, version=1)

        from docx import Document
        doc = Document(output)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Compliance Pilot" in all_text


class TestConfidenceTierMarkers:
    """Test that confidence tier markers are added to answers."""

    def test_needs_review_marker(self):
        qa = {"answer_text": "Some answer", "confidence_tier": "needs_review", "sources": []}
        result = _format_answer(qa)
        assert "\u26a0" in result
        assert "Needs Review" in result

    def test_no_answer_marker(self):
        qa = {"answer_text": "", "confidence_tier": "no_answer", "sources": []}
        result = _format_answer(qa)
        assert "\U0001f6a9" in result

    def test_auto_fill_has_answer(self):
        qa = {"answer_text": "Good answer", "confidence_tier": "auto_fill", "sources": []}
        result = _format_answer(qa)
        assert "Good answer" in result
        assert "\u26a0" not in result
        assert "\U0001f6a9" not in result

    @pytest.mark.asyncio
    async def test_markers_in_filled_docx(self, sample_docx, tmp_path):
        qa_pairs = _make_qa_pairs(
            answers=["Good answer", "Medium answer", ""],
            tiers=["auto_fill", "needs_review", "no_answer"],
        )

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_docx(str(sample_docx), qa_pairs, version=1)

        from docx import Document
        doc = Document(output)
        table = doc.tables[0]

        assert "\u26a0" in table.cell(2, 1).text
        assert "\U0001f6a9" in table.cell(3, 1).text
        assert "\u26a0" not in table.cell(1, 1).text
        assert "\U0001f6a9" not in table.cell(1, 1).text


# ---------------------------------------------------------------------------
# Failure cases
# ---------------------------------------------------------------------------


class TestFillPdfFormFields:
    """Test filling answers into PDF form fields."""

    @pytest.mark.asyncio
    async def test_fill_pdf_form_fields(self, sample_fillable_pdf, tmp_path):
        qa_pairs = [
            {
                "question_text": "What is your data retention policy?",
                "answer_text": "Data retained for 7 years.",
                "confidence_score": 0.9,
                "confidence_tier": "auto_fill",
                "location_info": {
                    "type": "pdf_form_field",
                    "page_number": 1,
                    "field_name": "data_retention_policy",
                },
                "sources": [{"source_file": "policy.pdf", "page_number": 1}],
            },
            {
                "question_text": "How do you handle data breaches?",
                "answer_text": "Notify within 24 hours.",
                "confidence_score": 0.85,
                "confidence_tier": "auto_fill",
                "location_info": {
                    "type": "pdf_form_field",
                    "page_number": 1,
                    "field_name": "breach_handling",
                },
                "sources": [],
            },
        ]

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_pdf(str(sample_fillable_pdf), qa_pairs, version=1)

        # Should output a PDF (not DOCX fallback)
        assert output.endswith(".pdf")
        assert os.path.exists(output)

        # Verify fields were filled
        import fitz
        doc = fitz.open(output)
        filled_fields = {}
        for page in doc:
            for widget in page.widgets():
                if widget.field_value:
                    filled_fields[widget.field_name] = widget.field_value
        doc.close()

        assert "data_retention_policy" in filled_fields
        assert "7 years" in filled_fields["data_retention_policy"]
        assert "breach_handling" in filled_fields
        assert "24 hours" in filled_fields["breach_handling"]

    @pytest.mark.asyncio
    async def test_fill_non_fillable_pdf_falls_back_to_docx(self, sample_pdf, tmp_path):
        """Non-fillable PDF should fall back to DOCX output."""
        qa_pairs = [
            {
                "question_text": "What is your data retention policy?",
                "answer_text": "7 years.",
                "confidence_score": 0.9,
                "confidence_tier": "auto_fill",
                "location_info": {"type": "pdf", "page_number": 1},
                "sources": [],
            },
        ]

        with patch("filler.OUTPUT_DIR", tmp_path):
            output = _fill_pdf(str(sample_pdf), qa_pairs, version=1)

        # Should fall back to DOCX
        assert output.endswith(".docx")
        assert os.path.exists(output)


class TestFillUnsupportedFormat:
    """Test that unsupported format raises ValueError."""

    @pytest.mark.asyncio
    async def test_fill_unsupported_format(self, tmp_path):
        with pytest.raises(ValueError, match="Unsupported file type"):
            await fill_document(str(tmp_path / "file.xyz"), "xyz", [], 1)
